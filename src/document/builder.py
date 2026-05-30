"""
builder.py — Build a reformatted .docx from extracted content + styles.
LLM decides structure; this module physically constructs the file.
"""

import json
import logging
from pathlib import Path
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH

logger = logging.getLogger(__name__)

ALIGNMENT_MAP = {
    "LEFT": WD_ALIGN_PARAGRAPH.LEFT,
    "CENTER": WD_ALIGN_PARAGRAPH.CENTER,
    "RIGHT": WD_ALIGN_PARAGRAPH.RIGHT,
    "JUSTIFY": WD_ALIGN_PARAGRAPH.JUSTIFY,
}


def build_docx(
    structure: list[dict],
    styles: dict,
    output_path: str,
) -> str:
    """
    Build a .docx file from structured content and style metadata.

    Args:
        structure: List of content blocks from LLM. Each block:
            {
                "type": "heading" | "paragraph" | "table" | "bullet_list",
                "level": 1-6 (for headings),
                "text": "...",
                "rows": [[cell, cell], ...] (for tables),
                "items": ["item1", ...] (for bullet_list),
            }
        styles: Output from style_parser.extract_styles()
        output_path: Where to save the .docx

    Returns:
        Absolute path of saved file.
    """
    doc = Document()

    # Apply page margins
    margins = styles.get("page_margins", {})
    for section in doc.sections:
        section.top_margin = Inches(margins.get("top_inches", 1.0))
        section.bottom_margin = Inches(margins.get("bottom_inches", 1.0))
        section.left_margin = Inches(margins.get("left_inches", 1.0))
        section.right_margin = Inches(margins.get("right_inches", 1.0))

    default_font = styles.get("default_font", {})
    heading_styles = styles.get("heading_styles", {})

    for block in structure:
        block_type = block.get("type", "paragraph")

        if block_type == "heading":
            level = max(1, min(int(block.get("level", 1)), 6))
            p = doc.add_heading(block.get("text", ""), level=level)
            _apply_heading_style(p, level, heading_styles, default_font)

        elif block_type == "paragraph":
            p = doc.add_paragraph(block.get("text", ""))
            _apply_default_font(p, default_font)
            alignment = styles.get("normal_style", {}).get("alignment", "LEFT")
            p.paragraph_format.alignment = ALIGNMENT_MAP.get(alignment, WD_ALIGN_PARAGRAPH.LEFT)

        elif block_type == "bullet_list":
            for item in block.get("items", []):
                p = doc.add_paragraph(item, style="List Bullet")
                _apply_default_font(p, default_font)

        elif block_type == "numbered_list":
            for item in block.get("items", []):
                p = doc.add_paragraph(item, style="List Number")
                _apply_default_font(p, default_font)

        elif block_type == "table":
            rows = block.get("rows", [])
            if not rows:
                continue
            num_cols = max(len(r) for r in rows)
            table = doc.add_table(rows=len(rows), cols=num_cols)
            table_style = styles.get("table_style")
            if table_style:
                try:
                    table.style = table_style
                except Exception:
                    table.style = "Table Grid"
            else:
                table.style = "Table Grid"

            for r_idx, row_data in enumerate(rows):
                for c_idx, cell_text in enumerate(row_data):
                    if c_idx < num_cols:
                        table.cell(r_idx, c_idx).text = str(cell_text)

        elif block_type == "page_break":
            doc.add_page_break()

    out = Path(output_path)
    out.parent.mkdir(parents=True, exist_ok=True)
    doc.save(str(out))
    logger.info(f"[Builder] Saved docx → {out}")
    return str(out.resolve())


def _apply_default_font(paragraph, font_def: dict):
    for run in paragraph.runs:
        run.font.name = font_def.get("name", "Calibri")
        size = font_def.get("size_pt")
        if size:
            run.font.size = Pt(size)
        color = font_def.get("color_rgb")
        if color:
            try:
                r, g, b = int(color[0:2], 16), int(color[2:4], 16), int(color[4:6], 16)
                run.font.color.rgb = RGBColor(r, g, b)
            except Exception:
                pass


def _apply_heading_style(paragraph, level: int, heading_styles: dict, default_font: dict):
    style_key = f"Heading {level}"
    h_style = heading_styles.get(style_key, {})
    for run in paragraph.runs:
        run.font.name = h_style.get("font_name") or default_font.get("name", "Calibri")
        size = h_style.get("size_pt")
        if size:
            run.font.size = Pt(size)
        if h_style.get("bold") is not None:
            run.font.bold = h_style["bold"]
        color = h_style.get("color_rgb")
        if color:
            try:
                r, g, b = int(color[0:2], 16), int(color[2:4], 16), int(color[4:6], 16)
                run.font.color.rgb = RGBColor(r, g, b)
            except Exception:
                pass
