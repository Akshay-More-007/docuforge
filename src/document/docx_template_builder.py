# -*- coding: utf-8 -*-
"""
docx_template_builder.py — DOCX template-mode builder.

Mirrors the proven PPTX template-mode: open the SOURCE .docx and modify it
IN PLACE so every design element is preserved:
  - theme (theme1.xml: fonts + scheme colors)
  - cover page styling
  - existing tables and their colored fills
  - headers/footers, margins, default font

It then applies targeted, grounded edits:
  - spelling / consistency fixes (run-level, formatting preserved)
  - inserts/refreshes a RACI matrix table (activities x stakeholders)
  - inserts/refreshes a vertical process-flow table (stakeholder | step | ▼)

Brand colors are EXTRACTED from the source (dominant table-header fill +
theme accent), so the generated tables stay on-brand for ANY input document.
"""

import logging
from collections import Counter
from pathlib import Path

from docx import Document
from docx.shared import Pt, RGBColor, Emu
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

logger = logging.getLogger(__name__)

# ── Sensible fallback palette (used only if extraction finds nothing) ─────────
_FALLBACK_BRAND   = "A91228"   # header / primary
_FALLBACK_ACCENT  = "1F3864"   # R/A accent (navy)
_GRAY_C           = "7B7B7B"   # Consulted
_GRAY_I           = "4A4A4A"   # Informed
_WHITE            = "FFFFFF"

# RACI code → fill color. R/A uses accent; C/I use grays; blank/'-' stays white.
def _raci_palette(brand: str, accent: str) -> dict:
    return {
        "R/A": accent, "A/R": accent,
        "R":   accent,
        "A":   brand,
        "C":   _GRAY_C,
        "I":   _GRAY_I,
        "-":   _WHITE, "":  _WHITE,
    }


# ── Public entry point ────────────────────────────────────────────────────────

def build_docx_from_template(
    source_path: str,
    output_path: str,
    *,
    spelling_fixes: dict | None = None,
    raci: dict | None = None,
    flow: dict | None = None,
    uniformity: dict | None = None,
    paragraph_edits: list | None = None,
    delete_paragraphs: list | None = None,
    table_edits: list | None = None,
    convert_to_table: list | None = None,
) -> str:
    """
    Open the source .docx, modify in place, save to output_path.

    Args:
        source_path: path to the source .docx (its theme/styling is preserved)
        output_path: where to save the result
        spelling_fixes: {wrong: right} run-level text replacements
        raci: {
            "activities": ["1. Vendor Empanelment", ...],
            "stakeholders": ["SRM Team", "Procurement", ...],
            "grid": [["R/A","C",...], ...]   # one row per activity, one col per stakeholder
        }
        flow: {
            "title": "START – Vendor Empanelment Request",
            "steps": [{"stakeholder": "SRM Team", "description": "Initial Evaluation ..."}, ...]
        }
        uniformity: {
            "apply": True,
            "font": "Arial" | None,        # None → keep the document's base font
            "header_fill": "C00000" | None # None → keep the extracted brand color
        }
        Normalizes every run to one font, repaints all dark table-header fills
        with the brand color, and unifies heading colors per level.
        paragraph_edits: [{"find": "<paragraph text>", "replace": "<new text>"}]
            Full-paragraph rewrite (style + first-run formatting preserved).
        delete_paragraphs: ["<exact paragraph text>", ...]
            Removes EVERY paragraph with that text — body line AND any manual
            TOC entry that duplicates it.
        table_edits: [{"match_header": "<text in first row>",
                       "headers": [...], "rows": [[...]]}        # full rebuild
                      | {"match_header": "...",
                         "cell_edits": [{"row": r, "col": c, "text": "..."}]}]
        convert_to_table: [{
            "after_heading": "<heading text>",
            "remove_until_heading": "<heading text>" | None,  # sweep stop (exclusive)
            "remove_anchor": False,        # also remove the anchor line itself
            "delete_headings": [...],      # extra texts to purge doc-wide (TOC sync)
            "intro": "<paragraph>" | None,
            "headers": [...], "rows": [[...]]
        }]

    Returns: absolute path to saved file.
    """
    doc = Document(source_path)

    brand, accent = _extract_brand_colors(doc)
    base_font = _extract_default_font(doc)
    logger.info(f"[DocxTemplate] brand={brand} accent={accent} font={base_font}")

    # Uniformity runs FIRST so the RACI/flow tables inserted below are styled
    # with the (possibly overridden) uniform font + brand and are never repainted.
    if uniformity and uniformity.get("apply"):
        base_font = uniformity.get("font") or base_font
        fill = uniformity.get("header_fill") or brand
        brand = fill.lstrip("#").upper()
        stats = _apply_uniformity(doc, font=base_font, header_fill=brand)
        logger.info(
            f"[DocxTemplate] Uniformity: font={base_font} fill={brand} "
            f"runs={stats['runs_refonted']} headers={stats['headers_recolored']} "
            f"headings={stats['headings_recolored']}"
        )

    # Structural ops next — they anchor on the ORIGINAL paragraph/table text,
    # so they must run before spelling fixes mutate that text.
    for op in (convert_to_table or []):
        try:
            ok = _convert_section_to_table(doc, op, brand, base_font)
            logger.info(f"[DocxTemplate] convert_to_table {op.get('after_heading')!r}: {ok}")
        except Exception as e:
            logger.error(f"[DocxTemplate] convert_to_table failed: {e}")

    for op in (table_edits or []):
        try:
            ok = _apply_table_edit(doc, op, brand, base_font)
            logger.info(f"[DocxTemplate] table_edit {op.get('match_header')!r}: {ok}")
        except Exception as e:
            logger.error(f"[DocxTemplate] table_edit failed: {e}")

    if paragraph_edits:
        n = _apply_paragraph_edits(doc, paragraph_edits)
        logger.info(f"[DocxTemplate] Applied {n} paragraph edit(s)")

    if delete_paragraphs:
        n = _delete_paragraphs(doc, delete_paragraphs)
        logger.info(f"[DocxTemplate] Deleted {n} paragraph(s)")

    if spelling_fixes:
        n = _apply_spelling_fixes(doc, spelling_fixes)
        logger.info(f"[DocxTemplate] Applied {n} spelling/consistency fix(es)")

    if raci and raci.get("activities") and raci.get("stakeholders"):
        ok = _insert_raci_after_heading(doc, raci, brand, accent, base_font)
        logger.info(f"[DocxTemplate] RACI matrix inserted={ok}")

    if flow and flow.get("steps"):
        ok = _insert_flow_after_heading(doc, flow, brand, accent, base_font)
        logger.info(f"[DocxTemplate] Process flow inserted={ok}")

    out = Path(output_path)
    out.parent.mkdir(parents=True, exist_ok=True)
    doc.save(str(out))
    logger.info(f"[DocxTemplate] Saved → {out}")
    return str(out.resolve())


# ── Brand / font extraction ─────────────────────────────────────────────────

def _extract_brand_colors(doc) -> tuple[str, str]:
    """
    Determine the source's brand (header) color and an accent color.

    brand  = most common non-white/non-auto cell fill across existing tables
    accent = theme dk2 (or a navy fallback), distinct from brand
    """
    header_fills: Counter = Counter()   # dark fills in first rows (true headers)
    dark_fills: Counter = Counter()     # dark fills anywhere (fallback)
    for t in doc.tables:
        for ri, row in enumerate(t.rows):
            for c in row.cells:
                f = _get_cell_fill(c)
                if not f:
                    continue
                fu = f.upper()
                if fu in ("FFFFFF", "AUTO", "FFFFF", "F2F2F2", "D9E1F2"):
                    continue
                if _luminance(fu) >= 0.6:
                    continue                # pastels/light grays are not brand
                dark_fills[fu] += 1
                if ri == 0:
                    header_fills[fu] += 1

    if header_fills:
        brand = header_fills.most_common(1)[0][0]
    elif dark_fills:
        brand = dark_fills.most_common(1)[0][0]
    else:
        brand = _FALLBACK_BRAND

    # Accent: prefer a dark navy from the theme color scheme; else fallback.
    accent = _theme_color(doc, "dk2") or _FALLBACK_ACCENT
    if accent.upper() == brand.upper():
        accent = _FALLBACK_ACCENT if brand.upper() != _FALLBACK_ACCENT else "0E2841"
    return brand, accent


def _theme_color(doc, key: str) -> str | None:
    """Read a theme scheme color (dk1/dk2/accent1...) from theme1.xml."""
    try:
        part = doc.part.package.part_related_by  # not used; read via zip below
    except Exception:
        pass
    try:
        import zipfile, re
        src = doc.part.package._partname  # unreliable; fall back to scanning parts
    except Exception:
        src = None
    # Most reliable: scan the theme part bytes already loaded in the package.
    try:
        for rel in doc.part.package.iter_parts():
            if rel.partname.endswith("theme1.xml") or "theme/theme" in str(rel.partname):
                xml = rel.blob.decode("utf-8", "ignore")
                import re
                m = re.search(rf"<a:{key}>\s*<a:srgbClr val=\"([0-9A-Fa-f]{{6}})\"", xml)
                if m:
                    return m.group(1).upper()
                m = re.search(rf"<a:{key}>\s*<a:sysClr[^>]*lastClr=\"([0-9A-Fa-f]{{6}})\"", xml)
                if m:
                    return m.group(1).upper()
    except Exception as e:
        logger.debug(f"[DocxTemplate] theme color read failed: {e}")
    return None


def _extract_default_font(doc) -> str:
    try:
        f = doc.styles["Normal"].font.name
        if f:
            return f
    except Exception:
        pass
    return "Arial"


# ── Spelling fixes (run-level, formatting preserved) ──────────────────────────

def _apply_spelling_fixes(doc, fixes: dict) -> int:
    count = 0

    # Safety rail: a case-only "fix" of an all-lowercase word (broadcasting →
    # Broadcasting) is a GLOBAL substring replace that corrupts mid-sentence
    # occurrences. Capitalisation of generic words belongs in paragraph_edits.
    fixes = {
        wrong: right for wrong, right in (fixes or {}).items()
        if wrong and right is not None and wrong != right
        and not (wrong.lower() == str(right).lower() and wrong == wrong.lower())
    }

    def fix_paragraph(p):
        nonlocal count
        for run in p.runs:
            txt = run.text
            if not txt:
                continue
            new = txt
            for wrong, right in fixes.items():
                if wrong and wrong in new:
                    new = new.replace(wrong, right)
            if new != txt:
                run.text = new
                count += 1

    for p in doc.paragraphs:
        fix_paragraph(p)
    for t in doc.tables:
        for row in t.rows:
            for c in row.cells:
                for p in c.paragraphs:
                    fix_paragraph(p)
    return count


# ── Granular structural edits (paragraphs, tables, list→table) ────────────────

def _norm(s: str) -> str:
    """Normalize text for matching: collapse whitespace, strip, casefold."""
    return " ".join(str(s or "").split()).casefold()


def _iter_body_paragraphs(doc):
    """Body paragraphs only (excludes table-cell paragraphs)."""
    from docx.text.paragraph import Paragraph
    for el in doc.element.body.iterchildren(qn("w:p")):
        yield Paragraph(el, doc)


def _find_paragraphs_by_text(doc, text: str) -> list:
    """All body paragraphs whose normalized text equals `text` (normalized)."""
    target = _norm(text)
    if not target:
        return []
    return [p for p in _iter_body_paragraphs(doc) if _norm(p.text) == target]


def _set_paragraph_text(p, new_text: str):
    """Replace a paragraph's text, keeping its style and first-run formatting."""
    runs = p.runs
    if runs:
        runs[0].text = new_text
        for r in runs[1:]:
            r._r.getparent().remove(r._r)
    else:
        p.add_run(new_text)


def _apply_paragraph_edits(doc, edits: list) -> int:
    """Full-paragraph rewrites. Exact-text match first, unique-substring fallback.

    Safety rail: a rewrite that shrinks the text below half its length is a
    destructive summary, not a fix — skipped unless the original contains a
    "(" annotation marker (annotation-stripping legitimately shortens text)."""
    count = 0
    for edit in edits or []:
        find = (edit or {}).get("find") or ""
        replace = (edit or {}).get("replace")
        if not find or replace is None or len(_norm(find)) < 4:
            continue
        if len(str(replace)) < 0.5 * len(find) and "(" not in find:
            logger.warning(f"[DocxTemplate] paragraph_edit skipped (shrinks too "
                           f"much): {find[:60]!r}")
            continue
        matches = _find_paragraphs_by_text(doc, find)
        if not matches:
            target = _norm(find)
            subs = [p for p in _iter_body_paragraphs(doc) if target in _norm(p.text)]
            # only act when unambiguous
            matches = subs if len(subs) == 1 else []
        for p in matches:
            _set_paragraph_text(p, replace)
            count += 1
    return count


def _delete_paragraphs(doc, texts: list) -> int:
    """Delete every body paragraph matching each text (body + manual-TOC copies).

    Safety rail: heading-styled paragraphs are never deleted here — real
    section headings may only be removed by a convert_to_table sweep (where
    they must be listed in delete_headings). Manual-TOC lines are Normal-styled
    so TOC cleanup still works."""
    count = 0
    for text in texts or []:
        if len(_norm(text)) < 4:        # guard against nuking empty/short paras
            continue
        for p in _find_paragraphs_by_text(doc, text):
            style = (p.style.name or "").lower() if p.style else ""
            if style.startswith(("heading", "title")):
                logger.warning(f"[DocxTemplate] delete skipped (heading): "
                               f"{text[:60]!r}")
                continue
            p._p.getparent().remove(p._p)
            count += 1
    return count


def _find_table_by_header(doc, match_header: str):
    """Find the table whose FIRST ROW contains match_header (normalized)."""
    target = _norm(match_header)
    if not target:
        return None
    for t in doc.tables:
        if not t.rows:
            continue
        first_row = " | ".join(c.text for c in t.rows[0].cells)
        if target in _norm(first_row):
            return t
    return None


def _content_col_widths(headers: list, rows: list, total_emu: int) -> list[int]:
    """Distribute total width across columns, weighted by content length.

    Short label columns (Sr. No etc.) stay narrow; description columns grow."""
    n = len(headers)
    weights = []
    for ci in range(n):
        longest = len(str(headers[ci]))
        for r in rows:
            if ci < len(r):
                longest = max(longest, len(str(r[ci])))
        # clamp: tiny columns get a floor, huge text saturates
        weights.append(max(6, min(longest, 90)))
    total_w = sum(weights)
    return [max(int(total_emu * w / total_w), 274320) for w in weights]  # ≥0.3"


def _body_content_width_emu(doc) -> int:
    """Usable text width of the first section (page width minus margins)."""
    try:
        sec = doc.sections[0]
        return int(sec.page_width - sec.left_margin - sec.right_margin)
    except Exception:
        return 5734050   # 6.27" fallback


def _build_styled_table(doc, headers: list, rows: list, brand: str, base_font: str):
    """Create a brand-styled table: brand header (white bold 10pt), 10pt body."""
    n_cols = max(len(headers), 1)
    tbl = doc.add_table(rows=len(rows) + 1, cols=n_cols)
    tbl.alignment = WD_TABLE_ALIGNMENT.CENTER

    for ci, h in enumerate(headers):
        _style_cell(tbl.rows[0].cells[ci], h, fill=brand, color=_WHITE,
                    bold=True, size=10, font=base_font)
    for ri, row in enumerate(rows, start=1):
        for ci in range(n_cols):
            txt = row[ci] if ci < len(row) else ""
            _style_cell(tbl.rows[ri].cells[ci], txt, fill=_WHITE, color="000000",
                        bold=False, size=10, font=base_font)

    _set_table_borders(tbl, "BFBFBF")
    widths = _content_col_widths(headers, rows, _body_content_width_emu(doc))
    for row in tbl.rows:
        for ci, w in enumerate(widths):
            row.cells[ci].width = Emu(w)
    return tbl


def _apply_table_edit(doc, op: dict, brand: str, base_font: str) -> bool:
    """Rebuild a table in place (headers+rows) or apply surgical cell edits."""
    tbl = _find_table_by_header(doc, op.get("match_header") or "")
    if tbl is None:
        logger.warning(f"[DocxTemplate] table_edit: no table matching "
                       f"{op.get('match_header')!r}")
        return False

    cell_edits = op.get("cell_edits")
    if cell_edits:
        for ce in cell_edits:
            try:
                r, c = int(ce.get("row", -1)), int(ce.get("col", -1))
                if 0 <= r < len(tbl.rows) and 0 <= c < len(tbl.rows[r].cells):
                    cell = tbl.rows[r].cells[c]
                    paras = cell.paragraphs
                    _set_paragraph_text(paras[0], str(ce.get("text", "")))
                    for extra in paras[1:]:
                        extra._p.getparent().remove(extra._p)
            except Exception as e:
                logger.warning(f"[DocxTemplate] cell_edit skipped: {e}")
        return True

    headers = op.get("headers") or []
    rows = op.get("rows") or []
    if not headers or not rows:
        return False
    new_tbl = _build_styled_table(doc, headers, rows, brand, base_font)
    tbl._tbl.addnext(new_tbl._tbl)
    tbl._tbl.getparent().remove(tbl._tbl)
    return True


def _convert_section_to_table(doc, op: dict, brand: str, base_font: str) -> bool:
    """Replace the free-text content of a section with a styled table.

    Sweeps blocks after `after_heading` until `remove_until_heading` (or the
    next heading if not given), inserts an optional intro paragraph + the new
    table, and purges `delete_headings` texts doc-wide (manual-TOC sync).
    """
    anchor_text = op.get("after_heading") or ""
    headers = op.get("headers") or []
    rows = op.get("rows") or []
    if not anchor_text or not headers or not rows:
        return False

    anchor = _find_heading(doc, [anchor_text]) if anchor_text else None
    if anchor is None:
        logger.warning(f"[DocxTemplate] convert_to_table: heading "
                       f"{anchor_text!r} not found")
        return False

    stop_text = _norm(op.get("remove_until_heading") or "")
    delete_headings = [t for t in (op.get("delete_headings") or []) if _norm(t)]
    sweepable = {_norm(t) for t in delete_headings}

    # ── Sweep: remove blocks between anchor and the stop heading ─────────────
    removed = 0
    nxt = anchor._p.getnext()
    while nxt is not None and removed < 100:
        tag = nxt.tag.split("}")[-1]
        if tag == "tbl":
            break                              # never sweep past an existing table
        if tag != "p":
            break                              # sectPr or other structural element
        p_text = _norm("".join(t.text or "" for t in nxt.iter(qn("w:t"))))
        if stop_text and p_text == stop_text:
            break
        # A heading always ends the sweep UNLESS it was explicitly merged away
        # (listed in delete_headings) — protects later sections when the LLM
        # supplies a wrong/never-matching stop heading.
        if _is_heading_p(nxt) and p_text and p_text not in sweepable:
            break
        following = nxt.getnext()
        nxt.getparent().remove(nxt)
        removed += 1
        nxt = following

    # ── Insert intro + table right after the anchor ───────────────────────────
    tbl = _build_styled_table(doc, headers, rows, brand, base_font)
    anchor._p.addnext(tbl._tbl)
    intro = op.get("intro")
    if intro:
        intro_el = OxmlElement("w:p")
        anchor._p.addnext(intro_el)
        from docx.text.paragraph import Paragraph
        ip = Paragraph(intro_el, anchor._parent)
        run = ip.add_run(str(intro))
        if base_font:
            run.font.name = base_font

    # ── Remove the anchor line itself when it is a stray annotation ──────────
    if op.get("remove_anchor"):
        anchor._p.getparent().remove(anchor._p)

    # ── TOC sync: purge any remaining copies of merged-away headings ─────────
    if delete_headings:
        _delete_paragraphs(doc, delete_headings)

    logger.info(f"[DocxTemplate] convert_to_table: swept {removed} block(s) "
                f"after {anchor_text!r}")
    return True


# ── Uniformity pass (fonts, table-header fills, heading colors) ───────────────

def _luminance(hex_color: str) -> float:
    """Perceived luminance 0..1; light pastels score high, brand darks low."""
    try:
        r, g, b = _hex_rgb(hex_color)
        return (0.299 * r + 0.587 * g + 0.114 * b) / 255.0
    except Exception:
        return 1.0


def _apply_uniformity(doc, *, font: str, header_fill: str) -> dict:
    """
    Normalize the document's look in three deterministic steps:
      1. fonts   — every run (body + tables) and every styled font set to `font`
      2. headers — first-row table cells with a DARK fill repainted `header_fill`
                   (light pastels like D9E1F2 are intentional sub-headers — kept;
                   flow tables are two-tone by design — skipped)
      3. headings — per heading level, recolor outlier runs to the level's
                   dominant explicit color
    """
    stats = {"runs_refonted": 0, "headers_recolored": 0, "headings_recolored": 0}

    # 1a. Style-level fonts: any style carrying an explicit font diverging from
    # the target gets rewritten, so style-inherited text follows too.
    for style in doc.styles:
        try:
            if style.font.name and style.font.name != font:
                style.font.name = font
        except Exception:
            continue
    try:
        doc.styles["Normal"].font.name = font
    except Exception:
        pass

    # 1b. Run-level fonts.
    def _refont(p):
        for run in p.runs:
            if run.text.strip() and run.font.name != font:
                run.font.name = font
                stats["runs_refonted"] += 1

    for p in doc.paragraphs:
        _refont(p)
    for t in doc.tables:
        for row in t.rows:
            for c in row.cells:
                for p in c.paragraphs:
                    _refont(p)

    # 2. Table header fills: repaint dark first-row fills to the brand color.
    for t in doc.tables:
        if _looks_like_flow(t._tbl):
            continue
        if not t.rows:
            continue
        for c in t.rows[0].cells:
            f = _get_cell_fill(c)
            if f and f.upper() != header_fill.upper() and _luminance(f) < 0.6:
                _set_cell_fill(c, header_fill)
                stats["headers_recolored"] += 1

    # 3. Heading colors: per level, find the dominant explicit run color and
    # apply it to every run of that level.
    by_level: dict[str, list] = {}
    for p in doc.paragraphs:
        name = p.style.name if p.style else ""
        if name.lower().startswith(("heading", "title")):
            by_level.setdefault(name, []).append(p)
    for name, paras in by_level.items():
        colors: Counter = Counter()
        for p in paras:
            for run in p.runs:
                if run.text.strip() and run.font.color and run.font.color.rgb:
                    colors[str(run.font.color.rgb)] += 1
        if not colors:
            continue
        dominant = colors.most_common(1)[0][0]
        rgb = RGBColor(*_hex_rgb(dominant))
        for p in paras:
            for run in p.runs:
                if not run.text.strip():
                    continue
                cur = str(run.font.color.rgb) if (run.font.color and run.font.color.rgb) else None
                if cur != dominant:
                    run.font.color.rgb = rgb
                    stats["headings_recolored"] += 1

    return stats


# ── RACI matrix insertion ─────────────────────────────────────────────────────

def _insert_raci_after_heading(doc, raci: dict, brand: str, accent: str, base_font: str) -> bool:
    activities   = raci["activities"]
    stakeholders = raci["stakeholders"]
    grid         = raci.get("grid") or []

    heading = _find_heading(doc, ["raci matrix", "raci"])
    if heading is None:
        heading = _append_heading(doc, "RACI Matrix")

    _remove_following_table(heading, matches=_looks_like_raci)

    n_rows = len(activities) + 1
    n_cols = len(stakeholders) + 1
    tbl = doc.add_table(rows=n_rows, cols=n_cols)
    tbl.alignment = WD_TABLE_ALIGNMENT.CENTER

    palette = _raci_palette(brand, accent)

    # Header row
    hdr = tbl.rows[0].cells
    _style_cell(hdr[0], "Activity", fill=brand, color=_WHITE, bold=True, size=9, font=base_font)
    for ci, sh in enumerate(stakeholders, start=1):
        _style_cell(hdr[ci], sh, fill=brand, color=_WHITE, bold=True, size=9, font=base_font,
                    align="center")

    # Body
    for ri, activity in enumerate(activities, start=1):
        cells = tbl.rows[ri].cells
        _style_cell(cells[0], activity, fill=_WHITE, color="000000", bold=False, size=9, font=base_font)
        row_codes = grid[ri - 1] if ri - 1 < len(grid) else []
        for ci in range(1, n_cols):
            code = (row_codes[ci - 1] if ci - 1 < len(row_codes) else "-") or "-"
            code = str(code).strip().upper().replace(" ", "")
            code = {"RA": "R/A", "AR": "A/R"}.get(code, code)
            fill = palette.get(code, _WHITE)
            is_filled = fill != _WHITE
            _style_cell(
                cells[ci], code if code not in ("", "-") else "-",
                fill=fill, color=_WHITE if is_filled else "000000",
                bold=is_filled, size=9, font=base_font, align="center",
            )

    _set_table_borders(tbl, "BFBFBF")
    _set_raci_col_widths(tbl, n_cols)
    _move_table_after(tbl, heading)
    _add_raci_legend(tbl, heading, palette, base_font)
    return True


def _set_raci_col_widths(tbl, n_cols):
    """First column wide (~1.75"), code columns narrow (~0.53")."""
    first = Emu(1600200)
    rest = Emu(482600)
    for row in tbl.rows:
        row.cells[0].width = first
        for ci in range(1, n_cols):
            row.cells[ci].width = rest


def _add_raci_legend(tbl, heading, palette, base_font):
    """Add a one-line legend paragraph right after the RACI table."""
    legend = OxmlElement("w:p")
    tbl._tbl.addnext(legend)
    from docx.text.paragraph import Paragraph
    p = Paragraph(legend, tbl._parent)
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    parts = [
        ("R = Responsible   ", "000000"),
        ("A = Accountable   ", "000000"),
        ("C = Consulted   ", "000000"),
        ("I = Informed   ", "000000"),
        ("R/A = Responsible & Accountable", "000000"),
    ]
    for text, col in parts:
        run = p.add_run(text)
        run.font.size = Pt(8)
        run.font.italic = True
        if base_font:
            run.font.name = base_font
        run.font.color.rgb = RGBColor(0x59, 0x59, 0x59)


# ── Process-flow insertion ─────────────────────────────────────────────────────

def _insert_flow_after_heading(doc, flow: dict, brand: str, accent: str, base_font: str) -> bool:
    steps = flow["steps"]
    title = flow.get("title", "START")

    heading = _find_heading(
        doc, ["end-to-end process flow", "end-to-end flow", "process flow", "process workflow"]
    )
    if heading is None:
        heading = _append_heading(doc, "Vendor Empanelment End-to-End Process Flow")

    _remove_following_table(heading, matches=_looks_like_flow)

    # Layout: header row + per step (content row + arrow row), minus trailing arrow
    n_rows = 1 + len(steps) * 2 - 1
    tbl = doc.add_table(rows=n_rows, cols=3)
    tbl.alignment = WD_TABLE_ALIGNMENT.CENTER

    # Header
    h = tbl.rows[0].cells
    _style_cell(h[0], "Stakeholder", fill=brand, color=_WHITE, bold=True, size=8, font=base_font, align="center")
    _style_cell(h[1], title, fill=accent, color=_WHITE, bold=True, size=9, font=base_font, align="center")
    _style_cell(h[2], "", fill=_WHITE, color="000000", bold=False, size=8, font=base_font)

    r = 1
    for i, step in enumerate(steps):
        stk = str(step.get("stakeholder", "")).strip()
        desc = str(step.get("description", "")).strip()
        cells = tbl.rows[r].cells
        _style_cell(cells[0], stk, fill=brand, color=_WHITE, bold=True, size=8, font=base_font, align="center")
        _style_cell(cells[1], desc, fill=accent if i == len(steps) - 1 else brand,
                    color=_WHITE, bold=True, size=9, font=base_font)
        _style_cell(cells[2], "", fill=_WHITE, color="000000", bold=False, size=8, font=base_font)
        r += 1
        # arrow row (skip after last step)
        if i < len(steps) - 1:
            acells = tbl.rows[r].cells
            _style_cell(acells[0], "", fill=_WHITE, color="000000", bold=False, size=8, font=base_font)
            _style_cell(acells[1], "▼", fill=_WHITE, color=brand, bold=True, size=12, font=base_font, align="center")
            _style_cell(acells[2], "", fill=_WHITE, color="000000", bold=False, size=8, font=base_font)
            r += 1

    _set_table_borders(tbl, "FFFFFF")  # flow uses fills, light/no visible borders
    _set_flow_col_widths(tbl)
    _move_table_after(tbl, heading)
    return True


def _set_flow_col_widths(tbl):
    w0, w1, w2 = Emu(914400), Emu(4114800), Emu(914400)
    for row in tbl.rows:
        row.cells[0].width = w0
        row.cells[1].width = w1
        row.cells[2].width = w2


# ── Low-level docx helpers ─────────────────────────────────────────────────────

def _style_cell(cell, text, *, fill, color, bold, size, font, align="left"):
    """Set a cell's text + fill + font in one shot, clearing prior content.

    Produces exactly ONE formatted run (no stray empty runs)."""
    _set_cell_fill(cell, fill)
    # remove extra paragraphs, keep the first
    while len(cell.paragraphs) > 1:
        pe = cell.paragraphs[-1]._p
        pe.getparent().remove(pe)
    p = cell.paragraphs[0]
    # remove all existing runs from the first paragraph
    for r in list(p.runs):
        r._r.getparent().remove(r._r)
    p.alignment = {
        "left": WD_ALIGN_PARAGRAPH.LEFT,
        "center": WD_ALIGN_PARAGRAPH.CENTER,
        "right": WD_ALIGN_PARAGRAPH.RIGHT,
    }.get(align, WD_ALIGN_PARAGRAPH.LEFT)
    run = p.add_run(str(text))
    run.font.size = Pt(size)
    run.font.bold = bold
    if font:
        run.font.name = font
    rgb = _hex_rgb(color)
    if rgb:
        run.font.color.rgb = RGBColor(*rgb)
    # tighten cell margins for compact tables
    _set_cell_margins(cell, top=40, bottom=40, left=60, right=60)


def _set_cell_fill(cell, hex_color: str):
    if not hex_color:
        return
    tcPr = cell._tc.get_or_add_tcPr()
    for shd in tcPr.findall(qn("w:shd")):
        tcPr.remove(shd)
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), hex_color.lstrip("#").upper())
    tcPr.append(shd)


def _set_cell_margins(cell, *, top, bottom, left, right):
    tcPr = cell._tc.get_or_add_tcPr()
    existing = tcPr.find(qn("w:tcMar"))
    if existing is not None:
        tcPr.remove(existing)
    tcMar = OxmlElement("w:tcMar")
    for tag, val in (("top", top), ("bottom", bottom), ("start", left), ("end", right),
                     ("left", left), ("right", right)):
        el = OxmlElement(f"w:{tag}")
        el.set(qn("w:w"), str(val))
        el.set(qn("w:type"), "dxa")
        tcMar.append(el)
    tcPr.append(tcMar)


def _set_table_borders(tbl, hex_color: str):
    tblPr = tbl._tbl.tblPr
    existing = tblPr.find(qn("w:tblBorders"))
    if existing is not None:
        tblPr.remove(existing)
    borders = OxmlElement("w:tblBorders")
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        el = OxmlElement(f"w:{edge}")
        el.set(qn("w:val"), "single")
        el.set(qn("w:sz"), "4")
        el.set(qn("w:space"), "0")
        el.set(qn("w:color"), hex_color.lstrip("#").upper())
        borders.append(el)
    tblPr.append(borders)


def _get_cell_fill(cell) -> str | None:
    tcPr = cell._tc.tcPr
    if tcPr is not None:
        shd = tcPr.find(qn("w:shd"))
        if shd is not None:
            f = shd.get(qn("w:fill"))
            if f and f != "auto":
                return f
    return None


def _find_heading(doc, keywords: list[str]):
    """
    Find the heading paragraph for a real body section matching any keyword.

    Documents often contain a manually-typed Table of Contents whose entries
    are ALSO styled as headings and match the same keywords. The TOC sits near
    the top, the real section near the bottom — so we return the LAST matching
    heading-styled paragraph (falling back to the last plain match).
    """
    kws = [k.lower() for k in keywords]
    last_heading = None
    last_any = None
    for p in doc.paragraphs:
        txt = p.text.strip().lower()
        if not txt or len(txt) >= 80:
            continue
        style = (p.style.name or "").lower() if p.style else ""
        is_heading = "heading" in style or "title" in style
        if any(k in txt for k in kws):
            last_any = p
            if is_heading:
                last_heading = p
    return last_heading or last_any


def _append_heading(doc, text: str):
    """Append a Heading-4-styled paragraph at the end (page break before)."""
    p = doc.add_paragraph()
    try:
        p.style = doc.styles["Heading 4"]
    except Exception:
        try:
            p.style = doc.styles["Heading 1"]
        except Exception:
            pass
    run = p.add_run(text)
    run.bold = True
    return p


def _is_heading_p(p_el) -> bool:
    """True if a <w:p> element is styled as a Heading/Title paragraph."""
    pPr = p_el.find(qn("w:pPr"))
    if pPr is None:
        return False
    pStyle = pPr.find(qn("w:pStyle"))
    if pStyle is None:
        return False
    val = (pStyle.get(qn("w:val")) or "").lower()
    return val.startswith("heading") or val.startswith("title")


def _table_text(tbl_el) -> str:
    """Concatenate all text in a <w:tbl> element (lowercased)."""
    texts = [t.text or "" for t in tbl_el.iter(qn("w:t"))]
    return " ".join(texts).lower()


def _looks_like_raci(tbl_el) -> bool:
    """A RACI grid: first header cell 'activity' OR many standalone R/A/C/I cells."""
    cells = [(t.text or "").strip() for t in tbl_el.iter(qn("w:t"))]
    if cells and cells[0].lower() in ("activity", "activities"):
        return True
    codes = sum(1 for c in cells if c.upper() in ("R", "A", "C", "I", "R/A", "A/R"))
    return codes >= 6


def _looks_like_flow(tbl_el) -> bool:
    """A process-flow table: contains ▼ arrows OR a 'stakeholder' header column."""
    txt = _table_text(tbl_el)
    if "▼" in txt:                      # ▼
        return True
    cells = [(t.text or "").strip().lower() for t in tbl_el.iter(qn("w:t"))]
    return bool(cells) and cells[0] == "stakeholder"


def _remove_following_table(heading_para, matches=None):
    """
    Remove the table(s) belonging to this heading's section.

    Scans forward to the next heading, but removes a <w:tbl> ONLY when the
    `matches(tbl_el)` predicate returns True. This guarantees we never delete an
    unrelated table (e.g. a definitions table) that happens to share a section.
    If `matches` is None, removes the first table found (legacy behavior).
    """
    nxt = heading_para._p.getnext()
    removed = False
    while nxt is not None:
        tag = nxt.tag.split("}")[-1]
        if tag == "tbl":
            following = nxt.getnext()
            if matches is None or matches(nxt):
                nxt.getparent().remove(nxt)
                removed = True
            nxt = following
            continue
        if tag == "p":
            if _is_heading_p(nxt):
                break          # reached the next section
            nxt = nxt.getnext()
            continue
        break                   # sectPr or other structural element
    return removed


def _move_table_after(tbl, heading_para):
    """Move a table element to sit right after the heading paragraph."""
    heading_para._p.addnext(tbl._tbl)


def _hex_rgb(hex_color: str):
    if not hex_color:
        return None
    h = str(hex_color).lstrip("#").strip()
    if len(h) != 6:
        return None
    try:
        return int(h[0:2], 16), int(h[2:4], 16), int(h[4:6], 16)
    except ValueError:
        return None
